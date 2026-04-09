from fastapi import FastAPI, Body
from fastapi.responses import JSONResponse, FileResponse
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile
import os
import re
import json
import time
import glob
from datetime import datetime
import traceback

app = FastAPI()
TEMP_DIR = tempfile.gettempdir()

# ... 排版辅助函数 ...
def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_three_line_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name, sz in [('top', '12'), ('bottom', '12')]: 
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    for border_name in ['left', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def set_header_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), 'auto')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

def add_separator_line(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def parse_content(paragraph, text, product_name=""):
    # 彻底过滤会导致 Word XML 崩溃的不可见控制字符
    text = "".join(c for c in str(text) if c.isprintable() or c in '\n\t')
    content = text.replace('\\n', '\n').replace('<br>', '\n').replace('<br/>', '\n')
    
    if product_name and product_name in content and f"**{product_name}**" not in content:
        content = content.replace(product_name, f"**{product_name}**")
    
    chunks = content.split('**')
    for i, chunk in enumerate(chunks):
        lines = chunk.split('\n')
        for j, line in enumerate(lines):
            if line:
                run = paragraph.add_run(line)
                if i % 2 == 1: run.bold = True
            if j < len(lines) - 1:
                paragraph.add_run().add_break()

@app.post("/generate_document")
async def generate_document(
    content: str = Body(...),
    product_name: str = Body(""),
    product_model: str = Body("")
):
    try:
        # 1. 垃圾清理防爆机制：清理 1 小时前的旧文档，防止 Render 硬盘塞满导致 500 崩溃
        now = time.time()
        for f in glob.glob(os.path.join(TEMP_DIR, "ETERNI_*.docx")):
            try:
                if os.stat(f).st_mtime < now - 3600:
                    os.remove(f)
            except: pass

        # 2. 残缺 JSON 抢救机制
        raw = content.strip().replace('```json', '').replace('```', '')
        parsed_successfully = False
        try:
            data = json.loads(raw)
            if isinstance(data, dict):
                for val in data.values():
                    if isinstance(val, str):
                        raw = val
                        parsed_successfully = True
                        break
        except: pass
        
        # 如果大模型超时导致 JSON 被截断（缺少后半个括号），强行用正则提取前面的正文
        if not parsed_successfully:
            match = re.search(r'"(?:docx|output|content)"\s*:\s*"(.*)', raw, re.DOTALL)
            if match:
                raw = match.group(1)
                raw = re.sub(r'\"\}?\s*$', '', raw)

        content = str(raw).replace('\\n', '\n').replace('\\"', '"')
        content = "".join(c for c in content if c.isprintable() or c in '\n\t')

        # 3. 完美的文档生成逻辑
        doc = Document()
        section = doc.sections[0]
        section.header_distance = Cm(0) 
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        fmt = style.paragraph_format
        fmt.line_spacing = 1.5
        fmt.space_before = fmt.space_after = Pt(0)

        header = section.header
        htable = header.add_table(1, 2, Inches(6.5))
        htable.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        htable.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        
        if os.path.exists("logo.png"):
            try: htable.rows[0].cells[0].paragraphs[0].add_run().add_picture("logo.png", width=Inches(0.6))
            except: pass
        
        p_meta = htable.rows[0].cells[1].paragraphs[0]
        p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_run = p_meta.add_run(f"{product_name}  {product_model}")
        meta_run.font.size = Pt(11)
        meta_run.font.color.rgb = RGBColor(120, 120, 120)
        add_separator_line(header.add_paragraph())

        lines = content.split('\n')
        in_table = False
        table_data = []
        
        first_p = doc.paragraphs[0]
        first_p_used = False

        for line in lines:
            stripped = line.strip()
            if not stripped: continue

            if stripped.startswith('|'):
                in_table = True
                if '---' not in stripped:
                    table_data.append([c.strip() for c in stripped.strip('|').split('|')])
                continue
            
            if in_table and not stripped.startswith('|'):
                if table_data and len(table_data) > 0:
                    first_p_used = True
                    max_cols = max(len(row) for row in table_data)
                    if max_cols > 0:
                        table = doc.add_table(rows=len(table_data), cols=max_cols)
                        set_three_line_borders(table)
                        for r_idx, row in enumerate(table_data):
                            for c_idx, val in enumerate(row):
                                if c_idx < max_cols: 
                                    cell = table.rows[r_idx].cells[c_idx]
                                    parse_content(cell.paragraphs[0], val, product_name)
                                    if r_idx == 0: set_header_border(cell)
                in_table = False; table_data = []

            if stripped.startswith('# '): 
                if not first_p_used:
                    p = first_p
                    first_p_used = True
                else:
                    p = doc.add_paragraph()
                    
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                title = stripped[2:]
                today = datetime.now().strftime("%B %d, %Y")
                
                run1 = p.add_run(title)
                run1.font.size = Pt(18); run1.bold = True
                p.add_run().add_break()
                
                run2 = p.add_run(f"ETERNI {product_name}")
                run2.font.size = Pt(15); run2.bold = True
                p.add_run().add_break()
                
                run3 = p.add_run(f"Issue Date: {today}")
                run3.font.size = Pt(15); run3.bold = True
                add_separator_line(doc.add_paragraph()) 

            elif stripped.startswith('## '): 
                if not first_p_used:
                    first_p_used = True
                else:
                    doc.add_paragraph() 
                add_separator_line(doc.add_paragraph()) 
                t_h2 = doc.add_table(1, 1) 
                cell = t_h2.rows[0].cells[0]
                set_cell_background(cell, '0033CC')
                run = cell.paragraphs[0].add_run(stripped[3:])
                run.font.size = Pt(14); run.font.color.rgb = RGBColor(255, 255, 255); run.bold = True
                doc.add_paragraph() 

            elif stripped.startswith('### '): 
                if not first_p_used:
                    p = first_p
                    first_p_used = True
                else:
                    doc.add_paragraph() 
                    p = doc.add_paragraph()
                    
                run = p.add_run(stripped[4:])
                run.font.size = Pt(12); run.bold = True

            else:
                if not first_p_used:
                    p = first_p
                    first_p_used = True
                else:
                    p = doc.add_paragraph()
                    
                if stripped.startswith(('- ', '* ')):
                    p.style = 'List Bullet'
                    parse_content(p, stripped[2:], product_name)
                else:
                    parse_content(p, stripped, product_name)

        if in_table and table_data and len(table_data) > 0:
            max_cols = max(len(row) for row in table_data)
            if max_cols > 0:
                table = doc.add_table(rows=len(table_data), cols=max_cols)
                set_three_line_borders(table)
                for r_idx, row in enumerate(table_data):
                    for c_idx, val in enumerate(row):
                        if c_idx < max_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            parse_content(cell.paragraphs[0], val, product_name)
                            if r_idx == 0: set_header_border(cell)

        file_name = f"ETERNI_{re.sub(r'[^a-zA-Z0-9]', '_', product_model)}.docx"
        path = os.path.join(TEMP_DIR, file_name)
        doc.save(path)
        return {"url": f"https://eterni-msds-api-1.onrender.com/download/{file_name}"}

    except Exception as e:
        # 全局异常捕获：如果再出问题，绝不报 500，而是把具体的 Python 错误直接打在屏幕上！
        error_trace = traceback.format_exc()
        print(error_trace) 
        return {"url": f"API内部崩溃，原因: {type(e).__name__} - {str(e)}"}

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(TEMP_DIR, filename)
    if os.path.exists(file_path):
        return FileResponse(file_path, filename=filename, headers={"Content-Disposition": f'attachment; filename="{filename}"'})
    return JSONResponse(status_code=404, content={"error": "File not found"})
